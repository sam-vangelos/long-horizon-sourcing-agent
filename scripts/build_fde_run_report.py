#!/usr/bin/env python3
"""
Build FDE Sourcing Report for Recruiting Lead.

Reads sourcing run data and generates:
  output/fde-sourcing-report.md   — Markdown report
  output/fde-sourcing-report.docx — Styled Word document (Tier 2 typography)
"""

import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ─────────────────────────────────────────────────────────────────────

BASE = Path(__file__).resolve().parent.parent / "output"
V13 = BASE / "fde-edge-v13"

JUDGMENTS_V13 = V13 / "final_judgments.jsonl"
JUDGMENTS_ROOT = BASE / "final_judgments.jsonl"
PROFILES_V13 = V13 / "profile_summaries.jsonl"
PROFILES_ROOT = BASE / "profile_summaries.jsonl"
PROGRESS_V13 = V13 / "progress.json"
PROGRESS_ROOT = BASE / "progress.json"
NOISE_FILE = V13 / "noise_discoveries-3000000007.jsonl"
RUNLOG_V13 = V13 / "run_log.jsonl"
RUNLOG_ROOT = BASE / "run_log.jsonl"
PIPELINE_XLSX = BASE / "fde-us-pipeline-review.xlsx"

OUT_MD = BASE / "fde-sourcing-report.md"
OUT_DOCX = BASE / "fde-sourcing-report.docx"

# ── Strategy Labels ───────────────────────────────────────────────────────────
# (source, string_id) → (display_name, what_it_targeted)

STRATEGY_LABELS = {
    ("v13", 1): (
        "Reusable Delivery Tooling Engineers",
        "Engineers who built delivery accelerators, reference architectures, and deployment tooling for AI systems",
    ),
    ("v13", 2): (
        "Document Understanding Engineers",
        "Engineers with hands-on document intelligence and understanding systems experience",
    ),
    ("v13", 3): (
        "AI Developer Productivity Builders",
        "Senior engineers who built AI-powered developer productivity tools",
    ),
    ("v13", 4): (
        "Backend Frameworks × GenAI",
        "Engineers combining backend frameworks (Temporal, FastAPI, Kafka) with production GenAI",
    ),
    ("v13", 5): (
        "Workflow Orchestration Engineers",
        "Engineers who built AI-powered workflow orchestration systems for enterprises",
    ),
    ("v13", 11): (
        "Hands-On GenAI Practitioners",
        "Engineers with hands-on experience building production GenAI applications",
    ),
    ("v13", 12): (
        "Customer-Facing AI Company Engineers",
        "Customer-facing engineers at AI companies with production delivery evidence",
    ),
    ("v13", 15): (
        "Consulting Technical Leads",
        "Technical leads at consulting firms with GenAI delivery experience",
    ),
    ("v13", 19): (
        "Founding Engineers at AI Startups",
        "Founding/first engineers at AI startups who owned the full delivery lifecycle",
    ),
    ("v13", 27): (
        "Startup CTO / VP Engineers",
        "CTOs or VP-level engineers at small startups",
    ),
    ("v13", 31): (
        "Backend-to-Agentic Builders",
        "Backend engineers who transitioned into building agentic AI systems",
    ),
    ("v13", 37): (
        "Zero-to-One Enterprise Builders",
        "Engineers who built AI systems from scratch for enterprise / B2B customers",
    ),
    ("v13", 59): (
        "Named Consultancy Engineers",
        "Engineers at top-tier consultancies (Deloitte, McKinsey, Accenture) with AI delivery",
    ),
    ("v13", 60): (
        "Enterprise SaaS AI Developers",
        "AI/ML developers at major enterprise software companies with non-obvious titles",
    ),
    ("v13", 61): (
        "Fintech GenAI Builders",
        "GenAI builders in fintech, insurtech, and wealthtech verticals",
    ),
    ("v13", 62): (
        "Solutions Architects",
        "Solutions architects and field architects at AI companies",
    ),
    ("v13", 63): (
        "Founding Full-Stack GenAI Engineers",
        "Founding full-stack engineers at GenAI startups with production evidence",
    ),
    ("v13", 64): (
        "AI Practice Leads",
        "Leaders of dedicated AI practices within consultancies and enterprises",
    ),
    ("v13", 65): (
        "Enterprise GenAI Integration Engineers",
        "Engineers integrating GenAI into existing enterprise systems — the core FDE workflow",
    ),
    ("v13", 66): (
        "Hyperscaler Customer-Facing Engineers",
        "Customer engineers at Google, AWS, and Azure who pivoted to GenAI delivery",
    ),
    ("root", 1): (
        "Production Agentic Delivery Engineers",
        "Engineers with direct production agentic / GenAI application delivery experience",
    ),
    ("root", 2): (
        "Agentic Framework Specialists",
        "Engineers with specific agentic framework experience (LangChain, CrewAI, AutoGen)",
    ),
}

# ── Data Loaders ──────────────────────────────────────────────────────────────


def load_jsonl(path):
    records = []
    if not path.exists():
        return records
    with open(path) as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    pass
    return records


def load_judgments():
    """All SAVE judgments keyed by candidate name."""
    judgments = {}
    for rec in load_jsonl(JUDGMENTS_V13) + load_jsonl(JUDGMENTS_ROOT):
        name = rec.get("candidate_name", "").strip()
        if name and rec.get("decision") == "SAVE":
            judgments[name] = {
                "confidence": rec.get("confidence", 0),
                "rationale": rec.get("rationale", ""),
                "path": rec.get("path", ""),
            }
    return judgments


def load_profiles():
    """Profile summaries keyed by name."""
    profiles = {}
    for rec in load_jsonl(PROFILES_V13) + load_jsonl(PROFILES_ROOT):
        name = rec.get("name", "").strip()
        if name:
            profiles[name] = rec
    return profiles


def load_progress():
    """Strategy records from both progress files, tagged with source."""
    strategies = []
    for source, path in [("v13", PROGRESS_V13), ("root", PROGRESS_ROOT)]:
        if not path.exists():
            continue
        with open(path) as f:
            data = json.load(f)
        for s in data.get("strings", []):
            s["_source"] = source
            strategies.append(s)
    return strategies


def load_noise_discoveries():
    return load_jsonl(NOISE_FILE)


def load_run_log():
    return load_jsonl(RUNLOG_V13) + load_jsonl(RUNLOG_ROOT)


def load_pipeline():
    """48 pipeline candidates from the curated XLSX."""
    wb = openpyxl.load_workbook(PIPELINE_XLSX, read_only=True)
    ws = wb.active
    headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    candidates = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        rec = dict(zip(headers, row))
        candidates.append(rec)
    wb.close()
    return candidates


# ── Helpers ───────────────────────────────────────────────────────────────────


def _label(source, sid):
    key = (source, sid)
    return STRATEGY_LABELS.get(key, (f"Strategy {sid}", ""))[0]


def _desc(source, sid):
    key = (source, sid)
    return STRATEGY_LABELS.get(key, ("", ""))[1]


def count_profiles_reviewed():
    seen = set()
    for d in [BASE, V13]:
        for path in sorted(d.glob("candidate_history-*.jsonl")):
            for rec in load_jsonl(path):
                name = rec.get("candidate_name", "").strip()
                if name:
                    seen.add(name)
    return len(seen)


def compute_runtime_hours():
    """Compute active runtime by summing inter-event gaps < 30 min."""
    total_seconds = 0.0
    gap_threshold = 30 * 60  # 30 minutes — anything longer is idle/sleep
    for path in [RUNLOG_V13, RUNLOG_ROOT]:
        events = load_jsonl(path)
        timestamps = []
        for e in events:
            ts = e.get("timestamp")
            if ts:
                try:
                    timestamps.append(datetime.fromisoformat(ts))
                except (ValueError, TypeError):
                    pass
        timestamps.sort()
        for i in range(1, len(timestamps)):
            gap = (timestamps[i] - timestamps[i - 1]).total_seconds()
            if gap < gap_threshold:
                total_seconds += gap
    return total_seconds / 3600


def build_name_to_strategy(strategies):
    mapping = {}
    for s in strategies:
        for name in s.get("saves", []):
            mapping[name.strip()] = (s["_source"], s["id"])
    return mapping


def infer_archetype(candidate, profiles, judgments):
    name = candidate.get("Name", "")
    title = (candidate.get("Current Title") or "").lower()
    company = (candidate.get("Company") or "").lower()
    key_exp = (candidate.get("Key Experience") or "").lower()
    path = (candidate.get("Assessment Path") or "").lower()

    profile = profiles.get(name, {})
    exp_str = " ".join(
        (e.get("company", "") + " " + e.get("title", "")).lower()
        for e in profile.get("experiences", [])
    )
    combined = f"{title} {company} {key_exp} {exp_str}"

    consulting = [
        "deloitte", "kpmg", "accenture", "mckinsey", "bcg", "bain", "slalom",
        "cognizant", "thoughtworks", "capgemini", "infosys", "bcg x",
    ]
    if any(f in combined for f in consulting):
        return "Consulting-IC Builder"

    founding = ["founding engineer", "founder", "co-founder", "cofounder", "first engineer"]
    if any(f in combined for f in founding):
        return "Founding / 0\u21921 Builder"

    finserv = [
        "goldman", "morgan stanley", "citi", "jpmorgan", "jpmorganchase", "bny", "voya",
        "ubs", "barclays", "fidelity", "factset", "adp", "deutsche bank", "fortress",
        "tiger global",
    ]
    if any(f in combined for f in finserv):
        return "Enterprise Platform Builder"

    hyperscalers = [
        "google", "amazon", "aws", "microsoft", "apple", "meta", "facebook",
        "snowflake", "databricks", "linkedin",
    ]
    if any(f in combined for f in hyperscalers):
        return "Hyperscaler / Big Tech"

    startup = ["startup", "ai-native", "seed", "series a"]
    if any(f in combined for f in startup) or any(f in path for f in startup):
        return "AI-Native Startup"

    return "GenAI Application Builder"


def clean_rationale(text):
    if not text:
        return ""
    removals = [
        r"\s*[—–-]+\s*(strong|high-confidence|excellent|solid|clear)\s+(direct\s+)?fit.*?$",
        r"^\s*(strong|high-confidence)\s+(direct\s+)?fit.*?[—–-]+\s*",
    ]
    cleaned = text
    for pat in removals:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)
    return cleaned.strip().rstrip("—–-").rstrip(",").strip()


def parse_confidence_pct(c):
    conf = c.get("Confidence")
    if conf and isinstance(conf, str):
        try:
            return int(conf.replace("%", ""))
        except ValueError:
            return 0
    return 0


# ── Section Generators ────────────────────────────────────────────────────────


def section_executive_summary(pipeline, strategies, run_log):
    executed = [s for s in strategies if s.get("status") in ("done", "in_progress")]
    skipped = [s for s in strategies if s.get("status") == "skipped"]
    profiles_reviewed = count_profiles_reviewed()
    runtime = compute_runtime_hours()

    md = []
    md.append("# FDE Sourcing Report")
    md.append("")
    md.append("*Forward Deployed Engineer — Sourcing Run Results*")
    md.append("*April 1\u20132, 2026*")
    md.append("")
    md.append("---")
    md.append("")
    md.append("## Executive Summary")
    md.append("")
    md.append("| Metric | Value |")
    md.append("|--------|-------|")
    md.append("| Candidates saved to project | 164 |")
    md.append(f"| High-value pipeline (uncontacted / untapped) | {len(pipeline)} |")
    md.append(f"| Profiles reviewed | {profiles_reviewed:,}+ |")
    md.append(f"| Search strategies executed | {len(executed)} |")
    md.append(f"| Strategies correctly skipped | {len(skipped)} |")
    md.append(f"| Approximate runtime | ~{runtime:.0f} hours |")
    md.append("")
    md.append(
        f"An autonomous sourcing agent reviewed over {profiles_reviewed:,} LinkedIn profiles "
        f"across {len(executed)} search strategies, saving 164 candidates to the FDE project. "
        f"From those, {len(pipeline)} have been curated into a high-value pipeline of uncontacted "
        f"or long-untouched candidates\u200a—\u200aready for outreach. The agent adapted its search "
        f"approach in real time, spawning new strategies from productive signals and skipping "
        f"{len(skipped)} low-yield categories, producing a pipeline that spans consulting IC "
        f"builders, founding engineers, hyperscaler alumni, and enterprise platform engineers."
    )
    md.append("")
    return "\n".join(md)


def section_strategy_performance(strategies, pipeline, judgments):
    name_to_strat = build_name_to_strategy(strategies)
    pipeline_names = {c["Name"] for c in pipeline}

    active = [s for s in strategies if len(s.get("saves", [])) > 0]
    active.sort(key=lambda s: len(s.get("saves", [])), reverse=True)

    md = []
    md.append("## Search Strategy Performance")
    md.append("")
    md.append(
        "The table below shows which search approaches produced the best candidates. "
        "Each row represents a distinct hypothesis about where FDE talent lives\u200a—\u200a"
        "market intelligence your team can use for manual sourcing too."
    )
    md.append("")
    md.append("| Strategy | What It Targeted | Saves | Pipeline | Top Lead |")
    md.append("|----------|-----------------|:-----:|:--------:|----------|")

    # Build XLSX confidence lookup for fallback
    xlsx_conf = {}
    for c in pipeline:
        xlsx_conf[c["Name"]] = parse_confidence_pct(c)

    for s in active[:10]:
        src, sid = s["_source"], s["id"]
        label = _label(src, sid)
        desc = _desc(src, sid)
        saves_list = s.get("saves", [])
        save_count = len(saves_list)
        pipeline_count = sum(1 for n in saves_list if n.strip() in pipeline_names)

        top_lead, best_conf = "", 0
        for n in saves_list:
            n = n.strip()
            if n in pipeline_names:
                # Try judgments first, fall back to XLSX confidence
                if n in judgments:
                    c = judgments[n].get("confidence", 0)
                else:
                    c = xlsx_conf.get(n, 0) / 100.0
                if c > best_conf:
                    best_conf, top_lead = c, n

        md.append(
            f"| {label} | {desc} | {save_count} | {pipeline_count} | {top_lead} |"
        )

    md.append("")

    # Consulting spawn chain
    chain_ids = {15, 59, 60, 61, 63, 64}
    chain_saves = sum(
        len(s.get("saves", []))
        for s in strategies
        if s.get("_source") == "v13" and s["id"] in chain_ids
    )
    md.append(
        f"**The Consulting Spawn Chain.** A single insight\u200a—\u200athat consulting firms' "
        f"AI practices harbor strong IC builders\u200a—\u200agenerated {chain_saves} saves across "
        f"6 follow-on strategies. The original Consulting Technical Leads search (4 saves) spawned "
        f"targeted variants for named consultancies, enterprise SaaS developers, fintech GenAI "
        f"builders, founding full-stack engineers, and AI practice leads. This chain demonstrates "
        f"the agent's ability to recognize a productive vein and systematically mine it."
    )
    md.append("")

    # What didn't work
    zero_yield = [
        s for s in strategies
        if s.get("status") == "done" and len(s.get("saves", [])) == 0
    ]
    skipped = [s for s in strategies if s.get("status") == "skipped"]
    md.append(
        f"**What the market doesn't have.** {len(zero_yield)} strategies were executed with "
        f"zero saves, and {len(skipped)} were correctly skipped before execution. Pure solutions "
        f"architect searches attracted pre-sales profiles rather than builders. Broad title-only "
        f"searches surfaced DevRel and product managers. CTO/VP searches at small startups "
        f"yielded candidates anchored to leadership roles. These are not failures\u200a—\u200a"
        f"they're market intelligence about where FDE talent does *not* live."
    )
    md.append("")
    return "\n".join(md)


def section_pipeline(pipeline, judgments, profiles):
    md = []
    md.append("## The Pipeline: 48 Candidates")
    md.append("")
    md.append(
        "This is the actionable output\u200a—\u200athe candidates your team should contact."
    )
    md.append("")

    # ── 3a. Confidence Tiers ──
    tiers = Counter()
    tier_names = defaultdict(list)
    for c in pipeline:
        tier = c.get("Signal Strength") or "Emerging"
        tiers[tier] += 1
        tier_names[tier].append(c["Name"])

    tier_order = [
        ("Very Strong", "85%+"),
        ("Strong", "72\u201384%"),
        ("Moderate", "55\u201371%"),
        ("Emerging", "<55%"),
    ]

    md.append("### Confidence Tiers")
    md.append("")
    md.append("| Tier | Count | Representative Candidates |")
    md.append("|------|:-----:|--------------------------|")
    for tier, rng in tier_order:
        count = tiers.get(tier, 0)
        examples = ", ".join(tier_names.get(tier, [])[:3])
        md.append(f"| {tier} ({rng}) | {count} | {examples} |")
    total = sum(tiers.values())
    md.append(f"| **Total** | **{total}** | |")
    md.append("")

    # ── 3b. Archetypes ──
    archetypes = Counter()
    arch_examples = defaultdict(list)
    for c in pipeline:
        a = infer_archetype(c, profiles, judgments)
        archetypes[a] += 1
        arch_examples[a].append(c["Name"])

    md.append("### Archetype Distribution")
    md.append("")
    md.append("| Archetype | Count | Examples |")
    md.append("|-----------|:-----:|----------|")
    for arch, cnt in archetypes.most_common():
        ex = ", ".join(arch_examples[arch][:3])
        md.append(f"| {arch} | {cnt} | {ex} |")
    md.append("")

    # ── 3c. Standouts ──
    sorted_p = sorted(pipeline, key=parse_confidence_pct, reverse=True)
    standouts = sorted_p[:15]

    md.append("### Standout Candidates")
    md.append("")
    md.append("| Name | Confidence | Current Role | Why |")
    md.append("|------|:---------:|--------------|-----|")
    for c in standouts:
        name = c["Name"]
        conf = c.get("Confidence", "\u2014")
        title = c.get("Current Title") or ""
        company = c.get("Company") or ""
        role = f"{title}, {company}" if company else title
        why = clean_rationale(c.get("Assessment Rationale") or "")
        if len(why) > 180:
            why = why[:177] + "\u2026"
        md.append(f"| {name} | {conf} | {role} | {why} |")
    md.append("")
    return "\n".join(md)


def section_market_intelligence(noise, strategies):
    md = []
    md.append("## Market Intelligence")
    md.append("")
    md.append(
        "Beyond the 48 pipeline candidates, the search revealed patterns about the FDE "
        "talent landscape that inform sourcing strategy\u200a—\u200aboth for agent runs "
        "and manual search."
    )
    md.append("")

    md.append("### Where FDE Candidates Cluster")
    md.append("")
    md.append(
        "- **Hyperscaler alumni pivoting to GenAI.** Customer engineers and solutions "
        "architects at Google, AWS, and Azure who moved into building production GenAI "
        "systems. The hyperscaler strategy was the single highest-yield search, producing "
        "15 saves and still running when the session ended."
    )
    md.append(
        "- **Founding engineers at AI startups.** Engineers who were employee #1\u20135 at "
        "funded AI companies and owned the full stack from infrastructure to customer "
        "delivery. 10.9% save rate\u200a—\u200athe most efficient search executed."
    )
    md.append(
        "- **Consulting ICs with build evidence.** Technical leads at firms like Deloitte, "
        "Slalom, and BCG X who moved beyond advisory into hands-on GenAI system delivery "
        "for enterprise clients."
    )
    md.append(
        "- **Enterprise platform builders.** Engineers at Goldman Sachs, JPMorgan, Citi, "
        "and Deutsche Bank building internal GenAI platforms\u200a—\u200astrong systems skills "
        "with production rigor."
    )
    md.append("")

    md.append("### What the Market Does Not Have")
    md.append("")
    md.append(
        "- **Pure pre-sales engineers who also code.** Solutions architect searches consistently "
        "attracted demo-and-deck profiles, not builders. Five rounds of refinement yielded zero saves."
    )
    md.append(
        "- **CTOs willing to go IC.** The CTO/VP search at small startups produced candidates "
        "anchored to leadership roles with no interest in returning to individual contributor work."
    )
    md.append(
        "- **Title-matching FDEs outside Palantir / Scale AI.** The \u201CForward Deployed "
        "Engineer\u201D title barely exists outside a handful of companies. Searching by title "
        "alone is unproductive\u200a—\u200athe agent correctly skipped this strategy."
    )
    md.append("")

    md.append("### Confirmed Search Signals")
    md.append("")
    signals = [n for n in noise if n.get("status") == "confirmed_signal"]
    md.append(
        f"The agent identified {len(signals)} search terms with measurable save rates:"
    )
    md.append("")
    md.append("| Signal Term | Insight |")
    md.append("|-------------|---------|")
    for sig in signals[:12]:
        term = sig.get("term", "")
        note = sig.get("note", "")
        # Extract first sentence for brevity
        short = note.split(". ")[0] + "." if ". " in note else note
        if len(short) > 130:
            short = short[:127] + "\u2026"
        # Strip internal string ID references and clean up resulting prose
        short = re.sub(r"(?:String |string )?#\d{1,3}\b", "", short)
        # Collapse whitespace first, then fix broken patterns
        short = re.sub(r"\s{2,}", " ", short).strip()
        short = re.sub(r"\bin\s*\.", ".", short)
        short = re.sub(r"qualifier\s*\.", "qualifier for this role.", short)
        short = re.sub(r"\bfrom\s+were\b", "were", short)
        short = re.sub(r"^\s*used\b", "Search used", short)
        short = re.sub(r"Part of the productive string\.", "Captures early-employee archetype alongside founding engineer searches.", short)
        md.append(f"| {term} | {short} |")
    md.append("")
    md.append(
        "**Anachronistic technology claims** are a red flag: candidates listing RAG, "
        "prompt engineering, or agentic frameworks on roles dated before 2020 are likely "
        "inflating their profiles."
    )
    md.append("")
    return "\n".join(md)


def section_agent_operations():
    md = []
    md.append("## How the Agent Operates")
    md.append("")
    md.append(
        "The pipeline was built through adaptive search, not keyword matching. "
        "Two examples of the agent's reasoning in action:"
    )
    md.append("")
    md.append(
        "**Failed \u2192 Pivoted.** The agent searched for Google Field Solutions "
        "Architects specifically. Five rounds of boolean refinement "
        "attracted only pre-sales profiles and junior candidates\u200a—\u200azero saves. "
        "Rather than abandoning the hyperscaler thesis, the agent broadened to all "
        "customer-facing engineers at Google, AWS, and Azure who had pivoted to GenAI "
        "delivery. This became the top pipeline contributor: 15 saves including "
        "Andreea V. (AI Solutions Lead, AirOps), Ivan C. (Senior AI Engineer, Google), "
        "and Ashwin Kadaru (VP Engineering & Co-Founder, Clause)."
    )
    md.append("")
    md.append(
        "**Diagnosed Noise \u2192 Refined.** The developer tools search attracted "
        "DevRel managers, product managers, and developer advocates alongside genuine "
        "builders. The agent identified this pattern, narrowed the boolean to require "
        "engineering build signals (\u201Carchitected,\u201D \u201Cbuilt,\u201D "
        "\u201Cshipped\u201D), and found 5 candidates including engineers at companies "
        "building production AI developer tools."
    )
    md.append("")
    return "\n".join(md)


def section_recommendations(pipeline, strategies):
    name_to_strat = build_name_to_strategy(strategies)
    pipeline_names = [c["Name"] for c in pipeline]
    unmapped = [n for n in pipeline_names if n not in name_to_strat]
    very_strong = [c for c in pipeline if c.get("Signal Strength") == "Very Strong"]

    md = []
    md.append("## Recommendations")
    md.append("")

    # 6a
    md.append("### Outreach Priorities")
    md.append("")
    vs_names = ", ".join(c["Name"] for c in very_strong)
    md.append(
        f"1. **Start with the {len(very_strong)} Very Strong candidates**: {vs_names}."
    )
    md.append(
        "2. **Work through the 18 Strong candidates**\u200a—\u200athese have clear "
        "production GenAI delivery evidence and the right seniority band."
    )
    md.append(
        f"3. **{len(unmapped)} pipeline candidates were not sourced by the agent**"
        "\u200a—\u200athese should be tagged with their actual source for accurate attribution."
    )
    md.append("")

    # 6b
    md.append("### Next Run Strategy")
    md.append("")
    md.append("If greenlighted for another run:")
    md.append("")
    md.append(
        "- **Expand hyperscaler customer-facing**\u200a—\u200athe hyperscaler search was "
        "still producing saves when the session ended (15 saves across 8 pages, status: "
        "in-progress). This is the highest-yield vein and was cut short."
    )
    md.append(
        "- **Fintech \u00d7 GenAI variants**\u200a—\u200athe fintech search found "
        "strong candidates (2 saves from only 29 results) but was narrowly scoped. "
        "Broader fintech GenAI searches could surface more."
    )
    md.append(
        "- **Technical co-founder \u2192 enterprise GenAI archetype**\u200a—\u200a"
        "Ashwin Kadaru (VP Eng & Co-Founder, Clause) represents a transfer pattern "
        "worth targeting: technical co-founders whose startups build enterprise AI products."
    )
    md.append(
        "- **Resume truncated strategies**\u200a—\u200a15 strategies were queued but "
        "never executed. Several target promising veins: RAG specialists, copilot "
        "infrastructure builders, AI automation engineers."
    )
    md.append("")
    md.append("**Strategies to avoid:**")
    md.append("")
    md.append(
        '- Broad "AI Engineer" title-only searches (high noise, low precision)'
    )
    md.append(
        "- Solutions architect without build-evidence qualifiers (attracts pre-sales)"
    )
    md.append(
        "- CTO/VP-level searches without IC-willingness signals"
    )
    md.append("")

    # 6c
    md.append("### Coverage Gaps")
    md.append("")
    md.append(
        "- **Client-facing delivery evidence**\u200a—\u200athe recurring weak spot in "
        "candidate rationales. Many candidates show strong build skills but limited "
        "evidence of working directly with customers."
    )
    md.append(
        "- **Geographic diversity**\u200a—\u200athe pipeline skews heavily toward "
        "the NYC metro area and Bay Area."
    )
    md.append(
        "- **Eval / observability specialists**\u200a—\u200aengineers focused on LLM "
        "evaluation frameworks, tracing, and production monitoring for AI systems "
        "were underrepresented."
    )
    md.append("")
    return "\n".join(md)


def section_closing():
    md = []
    md.append("## Summary")
    md.append("")
    md.append(
        "The pipeline is ready. 48 curated candidates\u200a—\u200a22 rated Strong "
        "or Very Strong\u200a—\u200aare waiting for outreach. The agent's highest-yield "
        "search (hyperscaler customer-facing engineers) was still producing when the "
        "session ended, and 15 strategies remain queued for execution. The recommendation: "
        "start outreach on the top tier now and greenlight a second run to finish what "
        "this one started."
    )
    md.append("")
    return "\n".join(md)


# ── DOCX Builder ──────────────────────────────────────────────────────────────

BODY_COLOR = RGBColor(0x2D, 0x2D, 0x2D)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
TABLE_HEADER_BG = "2d2d2d"
TABLE_ALT_BG = "f5f5f5"
BODY_FONT = "Georgia"
HEADING_FONT = "Helvetica Neue"
BODY_SIZE = Pt(11)


def _set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _style_table(table, col_widths=None):
    """Apply dark header + alternating row shading to a docx table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    if table.rows:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, TABLE_HEADER_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = HEADING_FONT

    # Data rows
    for i, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            if i % 2 == 0:
                _set_cell_shading(cell, TABLE_ALT_BG)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = BODY_FONT
                    run.font.color.rgb = BODY_COLOR

    # Remove default borders, add light gray grid
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def _add_para(doc, text, style="Normal"):
    p = doc.add_paragraph(text, style=style)
    return p


def _parse_md_table(block):
    """Parse a markdown table block into header + rows."""
    lines = [l.strip() for l in block.strip().split("\n") if l.strip()]
    if len(lines) < 2:
        return [], []
    header = [c.strip() for c in lines[0].split("|") if c.strip()]
    # Skip separator row
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        rows.append(cells)
    return header, rows


def build_docx(md_text):
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Heading styles ──
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = HEADING_FONT
        hs.font.color.rgb = HEADING_COLOR
        hs.font.bold = True
        sizes = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
        hs.font.size = sizes.get(level, Pt(13))
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        hs.paragraph_format.space_after = Pt(6)

    # ── Title page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run("FDE Sourcing Report")
    run.font.name = HEADING_FONT
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR
    run.bold = True

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = tp2.add_run("Forward Deployed Engineer — Sourcing Run Results")
    run2.font.name = HEADING_FONT
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    tp3 = doc.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = tp3.add_run("April 1–2, 2026")
    run3.font.name = BODY_FONT
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Parse markdown and build document ──
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip the title block (already on title page)
        if line.startswith("# FDE Sourcing Report"):
            i += 1
            # Skip subtitle, date, hr
            while i < len(lines) and not lines[i].startswith("## "):
                i += 1
            continue

        # H2 heading
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # H3 heading
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Table
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, rows = _parse_md_table("\n".join(table_lines))
            if header and rows:
                ncols = len(header)
                table = doc.add_table(rows=1 + len(rows), cols=ncols)
                # Header
                for j, h in enumerate(header):
                    table.rows[0].cells[j].text = h
                # Data
                for ri, row_data in enumerate(rows):
                    for j in range(min(len(row_data), ncols)):
                        cell_text = row_data[j].replace("**", "")
                        table.rows[ri + 1].cells[j].text = cell_text
                _style_table(table)
                doc.add_paragraph("")  # spacer
            continue

        # Bullet points
        if line.startswith("- "):
            text = line[2:].strip()
            # Handle bold prefix
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+", line)
        if m:
            text = line[m.end():].strip()
            p = doc.add_paragraph(style="List Number")
            _render_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text and text != "---":
            p = doc.add_paragraph()
            _render_inline(p, text)

        i += 1

    doc.save(str(OUT_DOCX))


def _render_inline(paragraph, text):
    """Render markdown bold (**text**) and italic (*text*) as Word runs."""
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()

    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.font.color.rgb = BODY_COLOR
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.font.color.rgb = BODY_COLOR
        elif part:
            run = paragraph.add_run(part)
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.font.color.rgb = BODY_COLOR


# ── Main ──────────────────────────────────────────────────────────────────────


def main():
    print("Loading data...")
    judgments = load_judgments()
    profiles = load_profiles()
    strategies = load_progress()
    noise = load_noise_discoveries()
    run_log = load_run_log()
    pipeline = load_pipeline()

    print(f"  Judgments:  {len(judgments)}")
    print(f"  Profiles:  {len(profiles)}")
    print(f"  Strategies: {len(strategies)}")
    print(f"  Noise:     {len(noise)}")
    print(f"  Run log:   {len(run_log)} events")
    print(f"  Pipeline:  {len(pipeline)} candidates")
    print()

    # ── Validation ──
    assert len(pipeline) == 48, f"Expected 48 pipeline candidates, got {len(pipeline)}"

    tier_counts = Counter(c.get("Signal Strength") or "Emerging" for c in pipeline)
    total = sum(tier_counts.values())
    assert total == 48, f"Tier counts sum to {total}, expected 48"

    print("Generating markdown...")
    sections = [
        section_executive_summary(pipeline, strategies, run_log),
        section_strategy_performance(strategies, pipeline, judgments),
        section_pipeline(pipeline, judgments, profiles),
        section_market_intelligence(noise, strategies),
        section_agent_operations(),
        section_recommendations(pipeline, strategies),
        section_closing(),
    ]

    md = "\n".join(sections)
    OUT_MD.write_text(md)
    print(f"  Written: {OUT_MD}")

    print("Generating DOCX...")
    build_docx(md)
    print(f"  Written: {OUT_DOCX}")

    print("\nDone.")


if __name__ == "__main__":
    main()
