"""Source-packet capture and document text extraction for intake authoring."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


SUPPORTED_SOURCE_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})
MAX_SOURCE_UPLOAD_BYTES = 5 * 1024 * 1024


class SourcePacketError(ValueError):
    """Raised when a source-packet upload cannot be read as text."""

    def __init__(self, message: str, *, code: str = "source_packet_error") -> None:
        super().__init__(message)
        self.code = code


@dataclass(frozen=True)
class ExtractedSourceFile:
    """One uploaded source artifact converted to recruiter-readable text."""

    filename: str
    content_type: str | None
    char_count: int
    text: str
    kind: str = "general"


def normalize_source_text(value: str | None) -> str:
    """Normalize pasted or extracted text without losing paragraph breaks."""

    if not value:
        return ""
    text = value.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def compose_source_packet_text(
    *,
    job_description_text: str = "",
    intake_notes_text: str = "",
    files: list[ExtractedSourceFile] | None = None,
    gap_answer_history: list[dict] | None = None,
) -> str:
    """Build the single synthesis input while preserving source boundaries."""

    sections: list[str] = []
    jd = normalize_source_text(job_description_text)
    notes = normalize_source_text(intake_notes_text)
    if jd:
        sections.append("JOB DESCRIPTION\n" + jd)
    if notes:
        sections.append("INTAKE NOTES\n" + notes)
    file_sections: dict[str, list[str]] = {
        "job_description": [],
        "intake_notes": [],
        "general": [],
    }
    for f in files or []:
        t = normalize_source_text(f.text)
        if t:
            kind = f.kind if f.kind in file_sections else "general"
            file_sections[kind].append(f"UPLOADED FILE: {f.filename}\n{t}")
    if file_sections["job_description"]:
        sections.append(
            "UPLOADED JOB DESCRIPTION FILES\n"
            + "\n\n".join(file_sections["job_description"])
        )
    if file_sections["intake_notes"]:
        sections.append(
            "UPLOADED INTAKE NOTES FILES\n"
            + "\n\n".join(file_sections["intake_notes"])
        )
    if file_sections["general"]:
        sections.append(
            "UPLOADED GENERAL FILES\n" + "\n\n".join(file_sections["general"])
        )
    answers = []
    for item in gap_answer_history or []:
        if not isinstance(item, dict):
            continue
        answer = normalize_source_text(str(item.get("answer_text") or ""))
        if answer:
            answers.append(answer)
    if answers:
        sections.append("RECRUITER GAP ANSWERS\n" + "\n\n".join(answers))
    return "\n\n---\n\n".join(sections).strip()


def extract_source_file_text(
    *,
    filename: str,
    content: bytes,
    content_type: str | None = None,
    kind: str = "general",
) -> ExtractedSourceFile:
    """Extract text from a supported uploaded source file."""

    if len(content) > MAX_SOURCE_UPLOAD_BYTES:
        raise SourcePacketError(
            f"{filename} is too large; max upload size is 5 MB.",
            code="source_file_too_large",
        )

    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_SOURCE_EXTENSIONS:
        raise SourcePacketError(
            f"{filename or 'file'} is not supported. Upload PDF, DOCX, TXT, or MD.",
            code="unsupported_source_file",
        )

    if suffix in {".txt", ".md"}:
        text = _decode_text(content)
    elif suffix == ".pdf":
        text = _extract_pdf_text(content, filename)
    else:
        text = _extract_docx_text(content, filename)

    text = normalize_source_text(text)
    if not text:
        raise SourcePacketError(
            f"{filename} did not contain readable text.",
            code="empty_source_file",
        )
    return ExtractedSourceFile(
        filename=filename,
        content_type=content_type,
        char_count=len(text),
        text=text,
        kind=_normalize_file_kind(kind),
    )


def _normalize_file_kind(kind: str | None) -> str:
    if kind in {"job_description", "intake_notes", "general"}:
        return kind
    return "general"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_pdf_text(content: bytes, filename: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - env dependency guard
        raise SourcePacketError(
            "PDF upload support requires pypdf to be installed.",
            code="source_parser_missing",
        ) from exc
    try:
        reader = PdfReader(BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 - parse libraries raise broad errors
        raise SourcePacketError(
            f"{filename} could not be parsed as a PDF.",
            code="source_parse_failed",
        ) from exc


def _extract_docx_text(content: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - env dependency guard
        raise SourcePacketError(
            "DOCX upload support requires python-docx to be installed.",
            code="source_parser_missing",
        ) from exc
    try:
        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if vals:
                    parts.append(" | ".join(vals))
        return "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise SourcePacketError(
            f"{filename} could not be parsed as a DOCX file.",
            code="source_parse_failed",
        ) from exc
